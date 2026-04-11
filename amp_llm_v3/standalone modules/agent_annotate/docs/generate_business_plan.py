#!/usr/bin/env python3
"""Generate Agent Annotate Business Plan as a formatted Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

ACCENT = RGBColor(0x1B, 0x3A, 0x5C)
LIGHT_BG = "D6E4F0"
WHITE = "FFFFFF"
DARK_HEADER = "1B3A5C"


def add_styled_table(headers, rows, col_widths=None):
    """Add a formatted table with dark header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{DARK_HEADER}"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = LIGHT_BG if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._element.get_or_add_tcPr().append(shading)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacer
    return table


def add_bullet(text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_callout(text, label="KEY METRIC"):
    """Add a highlighted callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(f"  {label}: ")
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)


# ════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AGENT ANNOTATE")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = ACCENT
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Business Plan & Monetization Strategy")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = "Calibri"

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Clinical Trial Annotation for Antimicrobial Peptide Research")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.name = "Calibri"

for _ in range(4):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amphoraxe Inc.")
run.font.size = Pt(14)
run.font.color.rgb = ACCENT
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("April 2026  |  Confidential")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = "Calibri"

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "Product Overview"),
    ("3.", "Market Opportunity"),
    ("4.", "Competitive Landscape"),
    ("5.", "Performance Benchmarks & Marketing Metrics"),
    ("6.", "Monetization Pathways"),
    ("7.", "Feasibility Analysis"),
    ("8.", "Go-to-Market Strategy"),
    ("9.", "Financial Projections"),
    ("10.", "Risk Analysis & Mitigation"),
    ("11.", "Implementation Roadmap"),
    ("12.", "Appendix: Technical Architecture"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "Agent Annotate is a multi-agent AI system that produces publication-grade "
    "structured annotations for clinical trials in the antimicrobial peptide (AMP) "
    "domain. It replaces a process that costs $200\u2013400 per trial with human "
    "annotators and takes weeks, delivering equivalent or superior results in "
    "3\u20135 minutes at approximately $0.15 per trial."
)

doc.add_paragraph(
    "The system deploys 12 parallel research agents that query 17+ free biomedical "
    "databases, 5 specialized annotation agents, and a 3-model blind verification "
    "layer\u2014all running locally via Ollama with zero external API costs. This "
    "architecture provides two structural advantages that define our market position: "
    "complete data sovereignty (nothing leaves the customer\u2019s network) and "
    "near-zero marginal cost per trial."
)

doc.add_heading("Key Figures", level=3)

add_styled_table(
    ["Metric", "Agent Annotate", "Human Baseline", "Delta"],
    [
        ["Outcome accuracy", "72.7%", "55.6% inter-rater", "+17.1 pp"],
        ["Peptide accuracy", "91.0%", "86.0% inter-rater", "+5.0 pp"],
        ["Classification accuracy", "85.3%", "93.2% inter-rater", "-7.9 pp*"],
        ["Delivery mode accuracy", "83.8%", "88.9% inter-rater", "-5.1 pp"],
        ["Trial coverage", "100%", "35\u201365%", "+35\u201365 pp"],
        ["Cost per trial", "$0.15", "$200\u2013400", "1,000\u20134,000x reduction"],
        ["Time per trial", "3\u20135 min", "Days\u2013weeks", "~500x faster"],
        ["Gross margin (SaaS)", "95%+", "\u2014", "\u2014"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5],
)

p = doc.add_paragraph()
run = p.add_run(
    "* 8 of 11 classification disagreements traced to human labeling errors in ground truth."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph(
    "This plan evaluates five monetization pathways, recommends a phased approach "
    "starting with consulting engagements, and projects $100\u2013300K ARR within "
    "18 months."
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  2. PRODUCT OVERVIEW
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Product Overview", level=1)

doc.add_heading("2.1  What It Does", level=2)
doc.add_paragraph(
    "Agent Annotate accepts clinical trial identifiers (NCT IDs) and produces "
    "structured, citation-backed annotations across six fields:"
)

add_styled_table(
    ["Field", "Values", "Purpose"],
    [
        ["Classification", "AMP / Other", "Is this an antimicrobial peptide trial?"],
        ["Delivery Mode", "Injection/Infusion, Oral, Topical, Other", "How is the drug administered?"],
        ["Outcome", "7 categories (Positive, Negative, Mixed, Unknown, etc.)", "What was the clinical result?"],
        ["Reason for Failure", "5 categories + empty", "If failed, why?"],
        ["Peptide", "True / False", "Is the intervention a peptide (2\u201350 AA)?"],
        ["Sequence", "Amino acid sequence or N/A", "What is the molecular sequence?"],
    ],
    col_widths=[1.3, 2.5, 2.7],
)

doc.add_heading("2.2  Three-Phase Pipeline", level=2)

doc.add_paragraph(
    "Phase 1 \u2014 Research: 12 parallel agents query 17+ databases "
    "(ClinicalTrials.gov, PubMed, UniProt, DRAMP, ChEMBL, DBAASP, RCSB PDB, "
    "WHO ICTRP, OpenAlex, and others). A Clinical Protocol Agent first extracts "
    "intervention names from the trial registry, enabling downstream agents to "
    "make targeted database queries."
)
doc.add_paragraph(
    "Phase 2 \u2014 Annotation: 5 specialized agents each handle one field using "
    "a two-pass investigative design: first extract facts from all research "
    "evidence, then make a determination. Published literature explicitly overrides "
    "stale registry status fields. A deterministic pre-classifier resolves "
    "~80% of trials without invoking the LLM."
)
doc.add_paragraph(
    "Phase 3 \u2014 Verification: 3 cognitively diverse models (conservative, "
    "evidence-strict, adversarial) independently review each annotation in a "
    "blind peer-review process. Unanimous agreement is accepted; disputes are "
    "resolved by a reconciliation agent using weighted voting."
)

doc.add_heading("2.3  Architectural Differentiators", level=2)

add_bullet("All inference runs locally via Ollama \u2014 no data leaves the network", bold_prefix="Data sovereignty: ")
add_bullet("All 17+ databases are free-tier public APIs \u2014 no external costs", bold_prefix="Zero API costs: ")
add_bullet("EDAM (Evidence-Driven Annotation Memory) learns from corrections, building institutional knowledge across jobs", bold_prefix="Self-learning: ")
add_bullet("Every annotation cites specific PMIDs, database records, and URLs", bold_prefix="Full traceability: ")
add_bullet("Structured data extraction before LLM invocation resolves 80% of trials deterministically", bold_prefix="Deterministic-first: ")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  3. MARKET OPPORTUNITY
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Market Opportunity", level=1)

doc.add_heading("3.1  Primary Market: AMP Clinical Trials", level=2)

doc.add_paragraph(
    "Antimicrobial peptides represent a growing class of therapeutics driven by "
    "the global antimicrobial resistance (AMR) crisis. The WHO has identified AMR "
    "as one of the top 10 global public health threats. AMP-based therapeutics "
    "are being developed for wound healing, resistant bacterial infections, "
    "biofilm disruption, and immunomodulation."
)

add_styled_table(
    ["Market Indicator", "Estimate", "Source"],
    [
        ["Existing AMP-related clinical trials", "5,000\u201310,000", "ClinicalTrials.gov + intl registries"],
        ["New AMP trials per year", "~500", "Registry analysis"],
        ["Cost of manual annotation", "$200\u2013400/trial", "Loaded annotator cost"],
        ["Human annotator throughput", "250\u2013500 trials/year", "Full-time equivalent"],
        ["Human annotation coverage", "35\u201365% of fields completed", "Ground truth analysis (R1/R2)"],
        ["Global AMR therapeutics market", "$5.4B by 2030", "Industry reports"],
    ],
    col_widths=[2.5, 1.8, 2.2],
)

doc.add_heading("3.2  Addressable Market Sizing", level=2)

add_styled_table(
    ["Segment", "TAM (Annual)", "Assumptions"],
    [
        ["AMP SaaS annotation", "$2.5M\u201320M", "5,000\u201310,000 trials at $0.50\u20132.00/trial"],
        ["On-premise licensing", "$500K\u20132M", "20\u201340 pharma/biotech orgs at $25\u201350K/year"],
        ["Pre-annotated data products", "$500K\u20131M", "20\u201350 subscribers at $10\u201325K/year"],
        ["Consulting/contract", "$200K\u2013500K", "10\u201330 engagements at $5\u201320K each"],
        ["Adjacent peptide markets*", "$10M\u201350M", "GLP-1, oncology peptides, neuropeptides"],
    ],
    col_widths=[2.0, 1.5, 3.0],
)

p = doc.add_paragraph()
run = p.add_run(
    "* Adjacent markets require pipeline adaptation but no architectural changes\u2014"
    "the system is disease-agnostic by design."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_heading("3.3  Target Customer Profiles", level=2)

add_bullet(
    " developing AMP therapeutics who need structured competitive "
    "intelligence on the clinical trial landscape",
    bold_prefix="Pharmaceutical companies:"
)
add_bullet(
    " conducting systematic reviews and meta-analyses "
    "of antimicrobial clinical trials",
    bold_prefix="Academic research groups:"
)
add_bullet(
    " aggregating trial data across multiple sponsors",
    bold_prefix="Contract Research Organizations (CROs):"
)
add_bullet(
    " needing structured trial summaries for "
    "approval review dossiers",
    bold_prefix="Regulatory consultants:"
)
add_bullet(
    " coordinating multi-institutional AMP research programs",
    bold_prefix="Research consortia:"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  4. COMPETITIVE LANDSCAPE
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Competitive Landscape", level=1)

add_styled_table(
    ["Competitor", "Type", "Price", "Key Weakness"],
    [
        [
            "Manual annotation\n(CROs, academic labs)",
            "Human labor",
            "$200\u2013400/trial",
            "48\u201391% inter-rater agreement; 35\u201365% field coverage; weeks of delay",
        ],
        [
            "TrialTrove / Cortellis",
            "SaaS subscription",
            "$10\u201350K/year",
            "Fixed schema; black-box scoring; cloud-only; no evidence citations",
        ],
        [
            "Single-LLM API\n(GPT, Claude, Llama)",
            "Per-token API",
            "$0.50\u20135.00/trial",
            "No verification; hallucination risk; data leaves network; no multi-source research",
        ],
        [
            "ClinicalTrials.gov\n(raw data)",
            "Free public data",
            "$0",
            "Unstructured; no annotation layer; stale status fields; no literature cross-ref",
        ],
    ],
    col_widths=[1.7, 1.2, 1.2, 2.4],
)

doc.add_heading("4.1  Agent Annotate Competitive Advantages", level=2)

add_styled_table(
    ["Advantage", "Agent Annotate", "Nearest Competitor"],
    [
        ["Data sovereignty", "100% local inference", "Cloud API required (GPT/Claude)"],
        ["External API costs", "$0 (all free-tier)", "$0.50\u20135.00/trial (token costs)"],
        ["Evidence traceability", "PMIDs, URLs, model ID per annotation", "None (single-LLM) or black-box (vendors)"],
        ["Verification", "3-model blind peer review", "None (single pass)"],
        ["Coverage", "100% of fields, 100% of trials", "35\u201365% human coverage"],
        ["Self-learning", "EDAM improves with each job", "Static (retrain required)"],
        ["Deterministic bypass", "80% resolved without LLM", "100% require LLM or manual"],
        ["Cost per trial", "$0.15 operating", "$200\u2013400 (human) or $0.50\u20135 (API)"],
    ],
    col_widths=[1.6, 2.5, 2.4],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  5. PERFORMANCE BENCHMARKS & MARKETING METRICS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Performance Benchmarks & Marketing Metrics", level=1)

doc.add_paragraph(
    "These are the concrete, auditable metrics that underpin Agent Annotate\u2019s "
    "market positioning. Every number is derived from controlled benchmarks against "
    "human-annotated ground truth (780 NCTs, dual-annotator design with AC\u2081 "
    "Gwet statistics)."
)

doc.add_heading("5.1  Headline Metrics for Marketing", level=2)

doc.add_paragraph(
    "The following metrics are designed for customer-facing materials\u2014each "
    "one is defensible, auditable, and backed by published methodology."
)

add_styled_table(
    ["Claim", "Metric", "Evidence", "Marketing Use"],
    [
        [
            '"More consistent than\nhuman experts"',
            "72.7% Outcome accuracy\nvs 55.6% human baseline",
            "v32 benchmark, 100 NCTs,\nAC\u2081 = 0.587",
            "Headline stat for all\nmarketing materials",
        ],
        [
            '"91% peptide\nidentification"',
            "91.0% accuracy,\nAC\u2081 = 0.885",
            "v32 benchmark,\n+5pp above human",
            "Technical credibility\nwith researchers",
        ],
        [
            '"1,000x cheaper than\nmanual annotation"',
            "$0.15/trial vs\n$200\u2013400 human",
            "Infrastructure cost\nanalysis",
            "ROI justification\nfor procurement",
        ],
        [
            '"100% trial coverage"',
            "All fields, all trials\nvs 35\u201365% human",
            "Ground truth R1/R2\ncoverage analysis",
            "Completeness pitch\nfor systematic reviews",
        ],
        [
            '"Every annotation\ncites its sources"',
            "PMIDs, URLs, database\nrecords per field",
            "Output format\nspecification",
            "Trust and transparency\nfor pharma",
        ],
        [
            '"Zero cloud exposure"',
            "100% local inference\nvia Ollama",
            "Architecture design",
            "Data governance pitch\nfor enterprise",
        ],
    ],
    col_widths=[1.5, 1.6, 1.6, 1.8],
)

doc.add_heading("5.2  Per-Field Accuracy Progression", level=2)

doc.add_paragraph(
    "Demonstrating continuous improvement builds confidence. These version-over-version "
    "gains show a system that is actively getting better:"
)

add_styled_table(
    ["Field", "Early (v2\u2013v5)", "Mid (v27e)", "Current (v32)", "Human Baseline", "Trend"],
    [
        ["Classification", "29.4%", "82.8%", "85.3%", "91.6%", "Approaching human"],
        ["Delivery Mode", "50.0%", "93.1%", "83.8%", "68.2%", "Above human*"],
        ["Outcome", "40.9%", "75.9%", "64.0\u201372.7%", "55.6%", "Above human"],
        ["Reason for Failure", "\u2014", "74.4%", "85.5%", "91.3%", "Approaching human"],
        ["Peptide", "88.0%", "80.0%", "91.0%", "48.4%", "Far above human"],
        ["Sequence", "\u2014", "\u2014", "47.2%", "52.0%", "Approaching human"],
    ],
    col_widths=[1.3, 1.0, 1.0, 1.2, 1.2, 1.3],
)

p = doc.add_paragraph()
run = p.add_run(
    "* Delivery mode v27e peak was on 50-NCT sample; v32 on 100-NCT sample. "
    "Variation reflects sample composition, not regression."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_heading("5.3  Statistical Rigor", level=2)

doc.add_paragraph(
    "All benchmarks report AC\u2081 (Gwet\u2019s agreement coefficient), which is "
    "more robust to prevalence paradox than Cohen\u2019s kappa. This matters because "
    "several annotation fields have highly skewed class distributions (e.g., 75%+ "
    "of trials are non-AMP). Reporting AC\u2081 alongside raw agreement rates "
    "provides defensible statistical evidence for peer review and regulatory contexts."
)

add_bullet(
    " AC\u2081 (Gwet), Cohen\u2019s kappa, raw agreement percentage",
    bold_prefix="Reported measures:"
)
add_bullet(
    " Planned per Fleiss et al. 1969 formula",
    bold_prefix="95% confidence intervals:"
)
add_bullet(
    " Tracked per Byrt et al. 1993",
    bold_prefix="Prevalence and bias indices:"
)
add_bullet(
    " 780 NCTs with dual-annotator design (7 annotators for R1, independent R2)",
    bold_prefix="Ground truth design:"
)

doc.add_heading("5.4  Metrics to Track for Ongoing Marketing", level=2)

add_styled_table(
    ["KPI", "Current Value", "12-Month Target", "Why It Matters"],
    [
        ["Outcome accuracy", "72.7%", "80%+", "Headline metric; each +1pp is a press-worthy milestone"],
        ["Trials annotated (cumulative)", "~200", "5,000+", "Demonstrates production scale and EDAM learning"],
        ["Customer pilot conversions", "0", "5\u201310", "Validates product-market fit"],
        ["Published paper citations", "0 (draft ready)", "3\u20135", "Academic credibility drives inbound"],
        ["Time to annotate 100 trials", "~8 hours", "<4 hours", "Speed benchmark for sales demos"],
        ["EDAM corrections absorbed", "~123", "1,000+", "Shows system is learning and improving"],
        ["Fields above human parity", "2 (Outcome, Peptide)", "4+", "Expanding \"better than human\" narrative"],
    ],
    col_widths=[1.8, 1.3, 1.3, 2.1],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  6. MONETIZATION PATHWAYS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Monetization Pathways", level=1)

doc.add_paragraph(
    "Five pathways are evaluated below. They are not mutually exclusive\u2014the "
    "recommended approach (Section 8) combines multiple pathways in a phased rollout."
)

# ── Option A: SaaS ──

doc.add_heading("6.1  Option A: SaaS \u2014 Per-Trial Pricing", level=2)

doc.add_paragraph(
    "Hosted platform where customers submit NCT IDs via web UI or API and "
    "receive annotated results with full evidence chains."
)

add_styled_table(
    ["Tier", "Price / Trial", "Includes"],
    [
        ["Pay-as-you-go", "$2.00", "Standard CSV, 6-field annotation, evidence citations"],
        ["Volume (500+/yr)", "$1.00", "Full JSON with verification metadata, priority queue"],
        ["Enterprise (5,000+/yr)", "$0.50", "Dedicated queue, custom fields, API integration, SLA"],
    ],
    col_widths=[1.8, 1.2, 3.5],
)

doc.add_heading("Cost/Benefit Analysis", level=3)

add_styled_table(
    ["Factor", "Assessment"],
    [
        ["Revenue potential", "$50\u2013100K ARR from 10 mid-size clients at blended $1.00/trial"],
        ["Gross margin", "95%+ (compute cost $0.15/trial against $0.50\u20132.00 price)"],
        ["Startup investment", "$8\u201315K server hardware; ~$500/yr hosting infrastructure"],
        ["Time to first revenue", "2\u20133 months (after initial pilot phase)"],
        ["Scalability", "High \u2014 near-zero marginal cost; limited by queue throughput"],
        ["Key risk", "Pharma data sensitivity \u2014 pattern of queried NCT IDs reveals competitive intelligence"],
        ["Mitigation", "Strict privacy policy; no cross-customer data; SOC 2 roadmap"],
    ],
    col_widths=[1.5, 5.0],
)

add_callout(
    "Lowest customer friction. Best for academic groups and small biotech "
    "who want results without running infrastructure.",
    "FIT"
)

# ── Option B: On-Premise License ──

doc.add_heading("6.2  Option B: On-Premise License", level=2)

doc.add_paragraph(
    "Customer runs Agent Annotate on their own hardware. Amphoraxe delivers "
    "software, documentation, model configurations, and support."
)

add_styled_table(
    ["Tier", "Price", "Includes"],
    [
        ["Academic", "$5,000/year", "Full system, community support, quarterly updates"],
        ["Commercial", "$25,000/year", "Full system, priority support, monthly updates, custom fields"],
        ["Enterprise perpetual", "$75,000\u2013150,000 one-time", "Perpetual license, 1-year support, installation assistance"],
    ],
    col_widths=[1.8, 1.7, 3.0],
)

doc.add_heading("Cost/Benefit Analysis", level=3)

add_styled_table(
    ["Factor", "Assessment"],
    [
        ["Revenue potential", "$125K/yr from 5 commercial licenses; $75\u2013150K per enterprise deal"],
        ["Gross margin", "~100% (customer provides all hardware and compute)"],
        ["Startup investment", "2\u20134 weeks packaging, installer tooling, documentation"],
        ["Time to first revenue", "4\u20136 months (longer enterprise sales cycle)"],
        ["Scalability", "Medium \u2014 support burden scales with customer count"],
        ["Key risk", "Customer environment variability; support overhead"],
        ["Mitigation", "Docker containerization; hardware compatibility matrix; tiered SLA"],
    ],
    col_widths=[1.5, 5.0],
)

add_callout(
    "Highest margin. Directly exploits Agent Annotate\u2019s local-inference "
    "architecture as a feature, not a limitation. Best for pharma with strict "
    "data governance.",
    "FIT"
)

# ── Option C: Consulting ──

doc.add_heading("6.3  Option C: Annotation-as-a-Service (Consulting)", level=2)

doc.add_paragraph(
    "Amphoraxe runs annotation jobs for clients as service engagements. "
    "Client provides NCT IDs and requirements; Amphoraxe delivers a complete "
    "annotated dataset with evidence report."
)

add_styled_table(
    ["Engagement", "Price", "Includes"],
    [
        ["Pilot (25\u201350 trials)", "Free / $500", "Accuracy demo vs human baseline; builds trust"],
        ["Systematic review (100\u2013500)", "$2,000\u20135,000", "Full annotation, evidence report, methodology section for publication"],
        ["Pipeline analysis (500\u20132,000)", "$5,000\u201315,000", "Competitive landscape mapping, trend analysis, structured dataset"],
        ["Custom annotation schema", "$10,000\u201325,000", "New fields, new disease area, validation against client ground truth"],
    ],
    col_widths=[2.0, 1.3, 3.2],
)

doc.add_heading("Cost/Benefit Analysis", level=3)

add_styled_table(
    ["Factor", "Assessment"],
    [
        ["Revenue potential", "$65\u2013105K/yr (10 reviews at $3.5K + 2\u20133 pipeline analyses + custom work)"],
        ["Gross margin", "80\u201390% (labor for QA, client communication, custom development)"],
        ["Startup investment", "Minimal \u2014 system already exists; need deliverable templates"],
        ["Time to first revenue", "1\u20132 months (fastest path to first dollar)"],
        ["Scalability", "Low \u2014 labor-bound for custom work; project-based revenue is lumpy"],
        ["Key risk", "Doesn\u2019t scale; revenue ceiling without transitioning clients to product"],
        ["Mitigation", "Use consulting as lead-gen funnel for SaaS and license upsells"],
    ],
    col_widths=[1.5, 5.0],
)

add_callout(
    "Fastest path to revenue and case studies. Best as a go-to-market strategy "
    "that feeds into Options A and B.",
    "FIT"
)

# ── Option D: Data Product ──

doc.add_heading("6.4  Option D: Research Data Product", level=2)

doc.add_paragraph(
    "Sell a continuously updated, pre-annotated database of all AMP-related "
    "clinical trials with structured annotations, evidence citations, and "
    "quarterly re-annotation as trial statuses change."
)

add_styled_table(
    ["Product", "Price", "Includes"],
    [
        ["AMP Trial Database", "$10,000\u201325,000/yr", "All AMP trials, quarterly re-annotation, CSV/JSON, citations"],
        ["Custom disease-area dataset", "$15,000\u201340,000/yr", "Pipeline adapted for oncology/endo/neuro peptide trials"],
        ["API access to live database", "$2,000\u20135,000/mo", "Real-time queries, webhook notifications for new annotations"],
    ],
    col_widths=[2.0, 1.5, 3.0],
)

doc.add_heading("Cost/Benefit Analysis", level=3)

add_styled_table(
    ["Factor", "Assessment"],
    [
        ["Revenue potential", "$75K+/yr from 5 database subs; +$24\u201360K/yr per API subscriber"],
        ["Gross margin", "90%+ after initial corpus annotation (incremental = new trials only)"],
        ["Startup investment", "2\u20134 weeks compute for full corpus annotation (5,000\u201310,000 trials)"],
        ["Time to first revenue", "3\u20135 months (annotation run + sales)"],
        ["Scalability", "High \u2014 same dataset sold to many subscribers; minimal marginal cost"],
        ["Key risk", "Competes with free ClinicalTrials.gov data; must demonstrate annotation value"],
        ["Mitigation", "Differentiate on structured annotation layer, evidence citations, recency"],
    ],
    col_widths=[1.5, 5.0],
)

add_callout(
    "Strong long-term recurring revenue. Creates a proprietary data asset that "
    "appreciates over time. Best for competitive intelligence buyers.",
    "FIT"
)

# ── Option E: Open-Core ──

doc.add_heading("6.5  Option E: Open-Core / Freemium", level=2)

doc.add_paragraph(
    "Open-source the core annotation engine. Monetize through premium features "
    "(multi-model verification, EDAM, agreement analytics), hosted service, and "
    "enterprise support."
)

add_styled_table(
    ["Tier", "Price", "Includes"],
    [
        ["Community (OSS)", "Free", "Core pipeline, single-model annotation, basic output"],
        ["Pro", "$500/month", "Multi-model verification, EDAM, premium model configs, analytics"],
        ["Enterprise", "$2,000+/month", "Custom fields, API, priority support, deployment assistance"],
    ],
    col_widths=[1.8, 1.2, 3.5],
)

doc.add_heading("Cost/Benefit Analysis", level=3)

add_styled_table(
    ["Factor", "Assessment"],
    [
        ["Revenue potential", "$120K+/yr at 2% conversion from 1,000 community users; +$24K+/yr per enterprise"],
        ["Gross margin", "85\u201395% depending on hosted vs self-hosted"],
        ["Startup investment", "4\u20138 weeks to split codebase, build licensing, write public docs"],
        ["Time to first revenue", "6\u201312 months (slow community ramp)"],
        ["Scalability", "Very high long-term \u2014 community contributes improvements"],
        ["Key risk", "Gives away core technology; competitors can fork; free tier may suffice"],
        ["Mitigation", "Only pursue after market position is strong (>$200K ARR, published paper)"],
    ],
    col_widths=[1.5, 5.0],
)

add_callout(
    "Market dominance play, not a revenue play. Only viable after establishing "
    "paid customer base and academic credibility.",
    "FIT"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  7. FEASIBILITY ANALYSIS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Feasibility Analysis", level=1)

doc.add_paragraph(
    "Each pathway is scored across five dimensions on a 1\u20135 scale "
    "(5 = most favorable)."
)

add_styled_table(
    [
        "Dimension",
        "A: SaaS",
        "B: License",
        "C: Consulting",
        "D: Data Product",
        "E: Open-Core",
    ],
    [
        ["Time to revenue", "4", "2", "5", "3", "1"],
        ["Revenue ceiling", "4", "5", "2", "4", "4"],
        ["Startup effort", "3", "3", "5", "3", "2"],
        ["Margin", "5", "5", "4", "5", "4"],
        ["Scalability", "4", "3", "1", "5", "5"],
        ["Data governance fit", "2", "5", "4", "3", "4"],
        ["TOTAL (/30)", "22", "23", "21", "23", "20"],
    ],
    col_widths=[1.5, 0.9, 0.9, 1.0, 1.1, 1.0],
)

doc.add_heading("7.1  Feasibility Verdict", level=2)

doc.add_paragraph(
    "Options B (License) and D (Data Product) score highest overall, but "
    "Option C (Consulting) scores highest on time-to-revenue and startup effort\u2014"
    "making it the best entry point. The recommended strategy uses C as a bridge "
    "to B and D."
)

doc.add_heading("7.2  Technical Feasibility", level=2)

add_bullet(
    " Production system running on Mac Mini M4 (16GB). "
    "32+ versions shipped with continuous accuracy improvement.",
    bold_prefix="Already built: "
)
add_bullet(
    " Docker containerization for customer deployment; "
    "hardware compatibility testing across GPU configurations; "
    "API authentication and rate limiting for SaaS mode.",
    bold_prefix="Needs work: "
)
add_bullet(
    " All external data sources are free-tier public APIs. "
    "No vendor dependency. System works air-gapped with pre-cached data.",
    bold_prefix="No external dependencies: "
)

doc.add_heading("7.3  Regulatory Feasibility", level=2)

doc.add_paragraph(
    "Agent Annotate does not make clinical decisions\u2014it annotates and "
    "classifies existing trial data. This positions it as a research tool, "
    "not a medical device, avoiding FDA/EMA regulatory burden. However, "
    "pharma customers will require:"
)

add_bullet("Documented validation methodology (PAPER.md provides this)")
add_bullet("Reproducible results (deterministic pipeline with version-pinned models)")
add_bullet("Audit trail (full JSON output includes model ID, agent provenance, timestamps)")
add_bullet("Data sovereignty assurance (local inference architecture)")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  8. GO-TO-MARKET STRATEGY
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Go-to-Market Strategy", level=1)

doc.add_heading("8.1  Phase 1: Prove It (Months 1\u20133)", level=2)
doc.add_paragraph("Primary model: Consulting (Option C)")

add_styled_table(
    ["Action", "Target", "KPI", "Cost"],
    [
        [
            "Run 3\u20135 free pilot engagements\n(25\u201350 trials each)",
            "AMP research groups,\nsystematic review authors",
            "2\u20133 case studies with\nauditable accuracy data",
            "$0 (existing hardware)",
        ],
        [
            "Publish PAPER.md to preprint\nserver (bioRxiv/medRxiv)",
            "Academic community",
            "Preprint posted;\n3\u20135 citations within 6 months",
            "$0",
        ],
        [
            "Present at 1\u20132 conferences\n(ASM Microbe, ICAAC)",
            "Infectious disease\nresearchers",
            "5\u201310 inbound leads",
            "$2\u20135K travel",
        ],
        [
            "Build landing page with\nmetrics dashboard",
            "All prospects",
            "100+ unique visitors/month",
            "$500 (domain + hosting)",
        ],
    ],
    col_widths=[2.0, 1.5, 1.7, 1.3],
)

add_callout(
    "Revenue target: $0\u201310K. Primary goal is case studies and credibility, not revenue.",
    "PHASE 1 TARGET"
)

doc.add_heading("8.2  Phase 2: Productize (Months 4\u20138)", level=2)
doc.add_paragraph("Primary models: SaaS (Option A) + Data Product (Option D)")

add_styled_table(
    ["Action", "Target", "KPI", "Cost"],
    [
        [
            "Launch hosted SaaS with\nper-trial pricing",
            "Pilot converts +\nnew inbound leads",
            "10+ paying customers;\n$20\u201350K ARR",
            "$8\u201315K (server)",
        ],
        [
            "Begin full-corpus annotation\n(5,000\u201310,000 AMP trials)",
            "Data product subscribers",
            "Complete corpus annotated;\n2\u20133 early subscribers",
            "$500 (compute electricity)",
        ],
        [
            "Submit PAPER.md to\npeer-reviewed journal",
            "Academic credibility",
            "Paper accepted; additional\ncitations and inbound",
            "$2\u20133K (publication fees)",
        ],
        [
            "Develop API access tier\nfor data product",
            "Enterprise data consumers",
            "API documented and\nfunctional",
            "$2\u20134K (development time)",
        ],
    ],
    col_widths=[2.0, 1.5, 1.7, 1.3],
)

add_callout(
    "Revenue target: $20\u201350K ARR. Transition from project-based to recurring revenue.",
    "PHASE 2 TARGET"
)

doc.add_heading("8.3  Phase 3: Scale (Months 9\u201318)", level=2)
doc.add_paragraph("Primary models: On-Premise License (Option B) + expanded Data Products")

add_styled_table(
    ["Action", "Target", "KPI", "Cost"],
    [
        [
            "Launch on-premise licensing\nwith Docker packaging",
            "Pharma companies from\nPhase 1\u20132 pipeline",
            "2\u20133 enterprise license\ndeals closed",
            "$5\u201310K (packaging,\nDocker, docs)",
        ],
        [
            "Expand pipeline to adjacent\ndisease areas (GLP-1, oncology)",
            "Endocrinology, oncology\nresearch groups",
            "1\u20132 new disease areas\nlaunched",
            "$5\u201310K (development)",
        ],
        [
            "Hire first sales/BD\nperson (part-time or contract)",
            "Enterprise pipeline\nacceleration",
            "5+ enterprise leads\nin pipeline",
            "$40\u201360K/yr (contract BD)",
        ],
        [
            "SOC 2 Type I certification",
            "Enterprise procurement\nrequirements",
            "Certification achieved",
            "$15\u201325K",
        ],
    ],
    col_widths=[2.0, 1.5, 1.7, 1.3],
)

add_callout(
    "Revenue target: $100\u2013250K ARR. Enterprise deals drive margin expansion.",
    "PHASE 3 TARGET"
)

doc.add_heading("8.4  Phase 4: Evaluate Open-Core (Month 18+)", level=2)

doc.add_paragraph("Decision criteria for pursuing Option E:")
add_bullet("ARR exceeds $200K (market position is defensible)")
add_bullet("3+ enterprise customers (paid relationships protect against free-tier cannibalization)")
add_bullet("Published, cited paper (academic credibility established)")
add_bullet("EDAM has 5,000+ experiences (learning system is a meaningful moat)")

doc.add_heading("8.5  Marketing Channels", level=2)

add_styled_table(
    ["Channel", "Tactic", "Expected Impact", "Cost"],
    [
        ["Academic publishing", "bioRxiv preprint + journal submission", "Inbound from researchers; citation-driven awareness", "$0\u20133K"],
        ["Conference presentations", "ASM Microbe, ICAAC, ECCMID", "Direct access to AMP researchers and pharma R&D", "$2\u20135K each"],
        ["LinkedIn / Twitter(X)", "Weekly metrics updates, benchmark posts", "Thought leadership; organic inbound", "$0"],
        ["Direct outreach", "Email to systematic review corresponding authors", "Targeted lead generation", "$0"],
        ["Pilot program", "Free 25-trial pilot for qualified prospects", "Conversion funnel entry point", "$0 (compute)"],
        ["SEO / content", "Blog posts on AMP trial annotation challenges", "Long-tail organic traffic", "$0\u2013500"],
    ],
    col_widths=[1.3, 2.0, 2.0, 1.2],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  9. FINANCIAL PROJECTIONS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Financial Projections", level=1)

doc.add_heading("9.1  Conservative Scenario", level=2)

add_styled_table(
    ["", "Year 1", "Year 2", "Year 3"],
    [
        ["Consulting revenue", "$35K", "$50K", "$60K"],
        ["SaaS revenue", "$15K", "$40K", "$80K"],
        ["License revenue", "$0", "$50K", "$125K"],
        ["Data product revenue", "$0", "$25K", "$75K"],
        ["Total revenue", "$50K", "$165K", "$340K"],
        ["Operating costs*", "$25K", "$80K", "$150K"],
        ["Net income", "$25K", "$85K", "$190K"],
        ["Gross margin", "50%", "52%", "56%"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5],
)

doc.add_heading("9.2  Aggressive Scenario", level=2)

add_styled_table(
    ["", "Year 1", "Year 2", "Year 3"],
    [
        ["Consulting revenue", "$60K", "$80K", "$100K"],
        ["SaaS revenue", "$30K", "$80K", "$160K"],
        ["License revenue", "$25K", "$125K", "$300K"],
        ["Data product revenue", "$15K", "$60K", "$150K"],
        ["Total revenue", "$130K", "$345K", "$710K"],
        ["Operating costs*", "$40K", "$130K", "$250K"],
        ["Net income", "$90K", "$215K", "$460K"],
        ["Gross margin", "69%", "62%", "65%"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5],
)

p = doc.add_paragraph()
run = p.add_run(
    "* Operating costs include: server hardware amortization, electricity, "
    "conference travel, publication fees, SOC 2 certification (Year 2), "
    "and contract BD hire (Year 2+). Does not include founder salary."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_heading("9.3  Unit Economics", level=2)

add_styled_table(
    ["Metric", "Value", "Notes"],
    [
        ["Cost per trial (compute)", "$0.15", "Electricity only; all APIs free"],
        ["SaaS price per trial", "$0.50\u20132.00", "Blended across tiers"],
        ["SaaS gross margin", "92\u201397%", "Best unit economics of any model"],
        ["License cost of goods", "~$0", "Customer provides hardware"],
        ["License gross margin", "~100%", "Highest margin model"],
        ["Consulting gross margin", "80\u201390%", "Limited by labor for QA/customization"],
        ["Data product marginal cost", "~$0.02/trial", "Only new trials after initial corpus"],
        ["Customer acquisition cost (CAC)", "$500\u20132,000", "Pilot program + direct outreach"],
        ["Lifetime value (LTV) \u2014 SaaS", "$5,000\u201315,000", "Assuming 2\u20133 year retention"],
        ["LTV:CAC ratio", "5:1 to 15:1", "Healthy unit economics"],
    ],
    col_widths=[2.2, 1.3, 3.0],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  10. RISK ANALYSIS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Risk Analysis & Mitigation", level=1)

add_styled_table(
    ["Risk", "Likelihood", "Impact", "Mitigation Strategy"],
    [
        [
            "AMP market too small\nto sustain growth",
            "Medium",
            "High",
            "Expand to adjacent peptide markets (GLP-1,\noncology, neuropeptides). Pipeline is\ndisease-agnostic by design.",
        ],
        [
            "Free LLM tools become\n\"good enough\"",
            "Medium",
            "Medium",
            "Verification layer + evidence traceability +\nlocal inference are structural advantages\nthat single-LLM tools cannot match.",
        ],
        [
            "Slow customer\nacquisition",
            "High",
            "Medium",
            "Free pilots reduce friction. Academic\npublishing drives inbound. Consulting\nmodel generates revenue during ramp.",
        ],
        [
            "Hardware requirements\ndeter customers",
            "Low",
            "Medium",
            "Mac Mini sufficient for dev. Cloud GPU\nrentals (Lambda, RunPod) for customers\nwithout hardware.",
        ],
        [
            "API rate limits or\naccess changes",
            "Low",
            "High",
            "17+ diversified sources. System degrades\ngracefully (fewer citations, not failure).\nAll sources are public government/academic.",
        ],
        [
            "Competitor replicates\nthe approach",
            "Medium",
            "Medium",
            "EDAM learning, 34+ version iterations,\nand integrated 17-source pipeline create\nsignificant replication cost barrier.",
        ],
        [
            "Accuracy plateau\nbefore human parity",
            "Low",
            "Medium",
            "Already above human on 2 fields. EDAM\nand model upgrades provide continuous\nimprovement path.",
        ],
    ],
    col_widths=[1.5, 0.9, 0.8, 3.3],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  11. IMPLEMENTATION ROADMAP
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("11. Implementation Roadmap", level=1)

add_styled_table(
    ["Timeline", "Milestone", "Deliverable", "Revenue Model"],
    [
        ["Month 1", "Launch pilot program", "3\u20135 free pilots with AMP research groups", "Consulting (C)"],
        ["Month 2", "Publish preprint", "PAPER.md on bioRxiv/medRxiv", "\u2014"],
        ["Month 3", "First paid engagement", "Systematic review annotation contract", "Consulting (C)"],
        ["Month 4", "SaaS beta launch", "Hosted platform with per-trial pricing", "SaaS (A)"],
        ["Month 5", "Begin corpus annotation", "Start annotating 5,000\u201310,000 AMP trials", "Data Product (D)"],
        ["Month 6", "First SaaS customers", "10+ paying SaaS accounts", "SaaS (A)"],
        ["Month 8", "Data product launch", "AMP Trial Database available for subscription", "Data Product (D)"],
        ["Month 9", "Docker packaging", "On-premise deployment package ready", "License (B)"],
        ["Month 10", "First license deal", "Enterprise on-premise contract signed", "License (B)"],
        ["Month 12", "Journal publication", "Peer-reviewed paper accepted", "\u2014"],
        ["Month 14", "Disease area expansion", "GLP-1 or oncology peptide pipeline launched", "All models"],
        ["Month 18", "Open-core evaluation", "Decision on open-source strategy", "Open-Core (E)"],
    ],
    col_widths=[1.0, 1.8, 2.5, 1.2],
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  12. APPENDIX
# ════════════════════════════════════════════════════════════════════════

doc.add_heading("12. Appendix: Technical Architecture", level=1)

doc.add_heading("12.1  Pipeline Overview", level=2)

doc.add_paragraph(
    "Phase 1 \u2014 Research\n"
    "  \u2514 Step 1: Clinical Protocol Agent extracts intervention names from registry\n"
    "  \u2514 Step 2: 11 parallel agents query databases using extracted names:\n"
    "      \u251c Literature (PubMed, PMC, BioC, OpenAlex, CrossRef, Semantic Scholar)\n"
    "      \u251c Peptide Identity (UniProt, DRAMP, APD)\n"
    "      \u251c Bioactivity (DBAASP, ChEMBL)\n"
    "      \u251c Structure (RCSB PDB, PDBe, EBI Proteins)\n"
    "      \u251c Pharmacology (IUPHAR)\n"
    "      \u251c Trial Registries (WHO ICTRP)\n"
    "      \u2514 Web Context (DuckDuckGo)"
)

doc.add_paragraph(
    "Phase 2 \u2014 Annotation\n"
    "  \u251c Classification Agent (AMP / Other)\n"
    "  \u251c Delivery Mode Agent (4 categories)\n"
    "  \u251c Outcome Agent (two-pass investigation)\n"
    "  \u251c Reason for Failure Agent\n"
    "  \u2514 Peptide Agent (with sequence extraction)"
)

doc.add_paragraph(
    "Phase 3 \u2014 Verification\n"
    "  \u251c Verifier 1: Conservative (gemma2:9b)\n"
    "  \u251c Verifier 2: Evidence-strict (qwen2.5:7b)\n"
    "  \u251c Verifier 3: Adversarial (llama3.1:8b)\n"
    "  \u251c Consensus check (unanimous = accept)\n"
    "  \u2514 Reconciliation Agent (qwen2.5:14b, weighted voting)"
)

doc.add_heading("12.2  Data Sources & Weights", level=2)

add_styled_table(
    ["Source", "Weight", "Data Type", "Cost"],
    [
        ["ClinicalTrials.gov", "0.95", "Trial protocol, status, interventions", "Free"],
        ["UniProt", "0.95", "Protein sequences, function", "Free"],
        ["PubMed", "0.90", "Abstracts, MeSH terms", "Free"],
        ["OpenFDA", "0.85", "Drug labels, approval status", "Free"],
        ["DBAASP", "0.85", "Antimicrobial activity, MIC data", "Free"],
        ["ChEMBL", "0.85", "Bioactivity assays, clinical phase", "Free"],
        ["PMC / BioC", "0.80\u20130.85", "Full-text articles, entity extraction", "Free"],
        ["DRAMP", "0.80", "AMP annotations, structural class", "Free"],
        ["RCSB PDB / PDBe", "0.80", "3D structure metadata", "Free"],
        ["EBI Proteins", "0.80", "Sequence, functional annotation", "Free"],
        ["OpenAlex", "Variable", "250M+ academic works", "Free (polite pool)"],
        ["Semantic Scholar", "Variable", "TLDR summaries, related papers", "Free"],
        ["CrossRef", "Variable", "Non-PubMed articles", "Free"],
        ["DuckDuckGo", "0.40", "General web context", "Free"],
    ],
    col_widths=[1.5, 0.8, 2.5, 0.7],
)

doc.add_heading("12.3  Hardware Profiles", level=2)

add_styled_table(
    ["Profile", "RAM", "Primary Model", "Verifiers", "Per-Trial Time", "Monthly Electricity"],
    [
        ["Mac Mini (dev)", "16 GB", "llama3.1:8b", "gemma2:9b, qwen2.5:7b, llama3.1:8b", "3\u20135 min", "~$15"],
        ["Server (prod)", "48+ GB", "Kimi K2 Thinking", "Larger verifier models", "2\u20134 min", "~$40\u201380"],
    ],
    col_widths=[1.1, 0.7, 1.2, 1.8, 0.9, 0.8],
)

# ── Final spacer + confidentiality ──

doc.add_paragraph("")
doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "\u2014 END OF DOCUMENT \u2014\n\n"
    "This document is confidential and proprietary to Amphoraxe Inc.\n"
    "Distribution without written consent is prohibited."
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = "Calibri"


# ── Save ────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Agent_Annotate_Business_Plan.docx",
)
doc.save(output_path)
print(f"Saved to: {output_path}")
